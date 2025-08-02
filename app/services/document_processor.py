from docx import Document
import spacy
import textstat
from typing import Dict, List, Any
from datetime import datetime
from app.services.debug_logger import DebugLogger

class DocumentProcessor:
    def __init__(self, app):
        self.app = app
        try:
            self.nlp = spacy.load("en_core_web_sm")
            DebugLogger.log_info("✅ spaCy model loaded successfully", {
                "model": "en_core_web_sm",
                "status": "ready"
            })
        except OSError:
            DebugLogger.log_error("spaCy model 'en_core_web_sm' not found. Please run: python -m spacy download en_core_web_sm")
            raise RuntimeError("spaCy model not loaded. Please install the required model.")
    
    def process_document(self, file_path: str, job_id: str) -> Dict[str, Any]:
        """Main document processing pipeline"""
        
        start_time = datetime.now()
        DebugLogger.log_document_processing(job_id, "started", {
            "file_path": file_path,
            "file_size": self._get_file_size(file_path)
        })
        
        try:
            # Update status
            self._update_status(job_id, "processing", "Starting analysis...")
            
            # Analyze paragraphs
            paragraph_analyses = self._analyze_paragraphs(file_path)
            
            # Generate overall analysis
            overall_analysis = self._generate_overall_analysis(paragraph_analyses)
            
            # Store results
            results = {
                'job_id': job_id,
                'file_path': file_path,
                'paragraph_analyses': paragraph_analyses,
                'overall_analysis': overall_analysis,
                'status': 'completed',
                'created_at': datetime.now().isoformat()
            }
            
            self.app.analysis_results[job_id] = results
            self._update_status(job_id, "completed", "Analysis completed")
            
            # Log performance
            duration = (datetime.now() - start_time).total_seconds()
            DebugLogger.log_performance("Document processing", duration, {
                "job_id": job_id,
                "paragraphs_analyzed": len(paragraph_analyses),
                "total_words": overall_analysis.get('total_words', 0)
            })
            
            DebugLogger.log_document_processing(job_id, "completed", {
                "duration": duration,
                "paragraphs": len(paragraph_analyses),
                "words": overall_analysis.get('total_words', 0)
            })
            
            return results
            
        except Exception as e:
            DebugLogger.log_error(f"Document processing failed for job {job_id}", e, {
                "file_path": file_path,
                "job_id": job_id
            })
            self._update_status(job_id, "failed", f"Analysis failed: {str(e)}")
            raise
    
    def _analyze_paragraphs(self, file_path: str) -> List[Dict[str, Any]]:
        """Analyze each paragraph"""
        
        try:
            doc = Document(file_path)
            analyses = []
            
            DebugLogger.log_debug("Starting paragraph analysis", {
                "file_path": file_path,
                "total_paragraphs": len(doc.paragraphs)
            })
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    analysis = self._analyze_single_paragraph(paragraph.text, i)
                    analyses.append(analysis)
            
            DebugLogger.log_info(f"Paragraph analysis completed", {
                "paragraphs_analyzed": len(analyses),
                "total_paragraphs": len(doc.paragraphs)
            })
            
            return analyses
        except Exception as e:
            DebugLogger.log_error(f"Error analyzing paragraphs", e, {
                "file_path": file_path
            })
            raise
    
    def _analyze_single_paragraph(self, text: str, index: int) -> Dict[str, Any]:
        """Analyze a single paragraph"""
        
        try:
            doc_nlp = self.nlp(text)
            
            analysis = {
                'paragraph_index': index,
                'text': text,
                'word_count': len(text.split()),
                'readability': textstat.flesch_reading_ease(text),
                'entities': [(ent.text, ent.label_) for ent in doc_nlp.ents],
                'comments': []
            }
            
            # Generate comments
            if analysis['word_count'] < 10:
                analysis['comments'].append("This paragraph is quite short. Consider adding more detail.")
            
            if analysis['readability'] < 30:
                analysis['comments'].append("This paragraph is difficult to read. Consider simplifying the language.")
            
            # Check for passive voice
            passive_count = sum(1 for token in doc_nlp if token.dep_ == "nsubjpass")
            if passive_count > 0:
                analysis['comments'].append(f"Consider using active voice instead of passive voice ({passive_count} instances).")
            
            DebugLogger.log_debug(f"Paragraph {index} analyzed", {
                "paragraph_index": index,
                "word_count": analysis['word_count'],
                "readability": analysis['readability'],
                "entities_found": len(analysis['entities']),
                "comments_generated": len(analysis['comments'])
            })
            
            return analysis
        except Exception as e:
            DebugLogger.log_error(f"Error analyzing paragraph {index}", e, {
                "paragraph_index": index,
                "text_length": len(text)
            })
            # Return basic analysis without NLP features
            return {
                'paragraph_index': index,
                'text': text,
                'word_count': len(text.split()),
                'readability': 0,
                'entities': [],
                'comments': ["Error analyzing this paragraph"]
            }
    
    def _generate_overall_analysis(self, paragraph_analyses: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate overall document analysis"""
        
        try:
            total_words = sum(p['word_count'] for p in paragraph_analyses)
            avg_readability = sum(p['readability'] for p in paragraph_analyses) / len(paragraph_analyses) if paragraph_analyses else 0
            total_entities = sum(len(p['entities']) for p in paragraph_analyses)
            
            analysis = {
                'total_paragraphs': len(paragraph_analyses),
                'total_words': total_words,
                'average_readability': avg_readability,
                'total_entities': total_entities,
                'overall_score': min(100, max(0, avg_readability))
            }
            
            DebugLogger.log_info("Overall analysis generated", {
                "total_paragraphs": analysis['total_paragraphs'],
                "total_words": analysis['total_words'],
                "average_readability": analysis['average_readability'],
                "overall_score": analysis['overall_score']
            })
            
            return analysis
        except Exception as e:
            DebugLogger.log_error(f"Error generating overall analysis", e)
            return {
                'total_paragraphs': len(paragraph_analyses),
                'total_words': 0,
                'average_readability': 0,
                'total_entities': 0,
                'overall_score': 0
            }
    
    def _update_status(self, job_id: str, status: str, message: str) -> None:
        """Update processing status"""
        
        self.app.processing_status[job_id] = {
            'status': status,
            'message': message,
            'timestamp': datetime.now().isoformat()
        }
        
        DebugLogger.log_debug(f"Status updated for job {job_id}", {
            "job_id": job_id,
            "status": status,
            "message": message
        })
    
    def _get_file_size(self, file_path: str) -> int:
        """Get file size in bytes"""
        try:
            import os
            return os.path.getsize(file_path)
        except Exception:
            return 0 